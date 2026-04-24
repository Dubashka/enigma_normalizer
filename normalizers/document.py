"""Нормализатор текстовых документов (TXT, DOCX).

Извлекает текст из документа, находит в нём PII-сущности (ФИО, адреса,
телефоны, ИНН, email, организации) и строит маппинг замен — аналогично
тому, как это делается для Excel-колонок.
"""
from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from normalizers.address import AddressNormalizer
from normalizers.email import EmailNormalizer
from normalizers.inn import InnNormalizer
from normalizers.organization import OrganizationNormalizer
from normalizers.person import PersonNormalizer
from normalizers.phone import PhoneNormalizer


# ---------------------------------------------------------------------------
# Детекторы сущностей (regex + вызов нормализаторов)
# ---------------------------------------------------------------------------

_PHONE_RE = re.compile(
    r"(?<!\d)"
    r"(\+7|8|7)?[\s\-\(]*\d{3}[\s\-\)]*\d{3}[\s\-]*\d{2}[\s\-]*\d{2}"
    r"(?!\d)"
)

_EMAIL_RE = re.compile(
    r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}"
)

_INN_RE = re.compile(r"(?<!\d)(\d{10}|\d{12})(?!\d)")

_FIO_RE = re.compile(
    r"\b([А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?)\s+"
    r"([А-ЯЁ][а-яё]+)\s+"
    r"([А-ЯЁ][а-яё]+)\b"
)

_FIO_INITIALS_RE = re.compile(
    r"\b([А-ЯЁ][а-яё]+)\s+([А-ЯЁ])\.\s*([А-ЯЁ])\."
)

_ADDRESS_KEYWORDS = re.compile(
    r"(?i)(ул\.|улица|пр\.|проспект|пер\.|переулок|бул\.|бульвар|"
    r"пл\.|площадь|ш\.|шоссе|наб\.|набережная|д\.\s*\d|кв\.\s*\d|"
    r"г\.\s+[А-ЯЁ]|обл\.|область|р-н|район|корп\.|строение|стр\.)"
)

_ORG_KEYWORDS = re.compile(
    r"(?i)\b(ООО|ОАО|ЗАО|ПАО|АО|ИП|НКО|АНО|ФГУП|МУП|ГУП|ГК|"
    r"LLC|OJSC|CJSC|JSC)\b"
)


@dataclass
class EntityMatch:
    """Найденная сущность в тексте."""
    entity_type: str        # fio / phone / email / inn / address / organization
    original: str           # оригинальное значение в тексте
    canonical: str          # нормализованное каноническое значение
    start: int              # позиция начала в тексте
    end: int                # позиция конца в тексте
    confidence: float = 1.0
    meta: dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentNormalizationResult:
    """Результат нормализации документа."""
    normalized_text: str
    mapping: dict[str, Any]
    entities: list[EntityMatch]
    stats: dict[str, int]


# ---------------------------------------------------------------------------
# Извлечение текста из файлов
# ---------------------------------------------------------------------------

def extract_text_from_txt(content: bytes, encoding: str = "utf-8") -> str:
    for enc in (encoding, "utf-8-sig", "cp1251", "latin-1"):
        try:
            return content.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


def extract_text_from_docx(content: bytes) -> tuple[str, Any]:
    """Возвращает (plain_text, document_object) для последующей замены."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError("Установите python-docx: pip install python-docx")
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs), doc


# ---------------------------------------------------------------------------
# Поиск сущностей
# ---------------------------------------------------------------------------

def _find_phones(text: str) -> list[EntityMatch]:
    normalizer = PhoneNormalizer()
    matches = []
    for m in _PHONE_RE.finditer(text):
        raw = m.group().strip()
        try:
            canonical = normalizer.normalize_value(raw)
        except Exception:
            canonical = raw
        matches.append(EntityMatch(
            entity_type="phone",
            original=raw,
            canonical=canonical,
            start=m.start(),
            end=m.end(),
            confidence=0.95,
        ))
    return matches


def _find_emails(text: str) -> list[EntityMatch]:
    normalizer = EmailNormalizer()
    matches = []
    for m in _EMAIL_RE.finditer(text):
        raw = m.group().strip()
        try:
            canonical = normalizer.normalize_value(raw)
        except Exception:
            canonical = raw
        matches.append(EntityMatch(
            entity_type="email",
            original=raw,
            canonical=canonical,
            start=m.start(),
            end=m.end(),
            confidence=0.99,
        ))
    return matches


def _find_inns(text: str) -> list[EntityMatch]:
    normalizer = InnNormalizer()
    matches = []
    for m in _INN_RE.finditer(text):
        raw = m.group().strip()
        try:
            canonical = normalizer.normalize_value(raw)
            if canonical == raw:
                confidence = 0.6
            else:
                confidence = 0.9
        except Exception:
            canonical = raw
            confidence = 0.5
        matches.append(EntityMatch(
            entity_type="inn",
            original=raw,
            canonical=canonical,
            start=m.start(),
            end=m.end(),
            confidence=confidence,
        ))
    return matches


def _find_fio(text: str) -> list[EntityMatch]:
    normalizer = PersonNormalizer()
    matches = []
    seen_spans: list[tuple[int, int]] = []

    for m in _FIO_RE.finditer(text):
        raw = m.group().strip()
        try:
            canonical = normalizer.normalize_value(raw)
        except Exception:
            canonical = raw
        matches.append(EntityMatch(
            entity_type="fio",
            original=raw,
            canonical=canonical,
            start=m.start(),
            end=m.end(),
            confidence=0.85,
        ))
        seen_spans.append((m.start(), m.end()))

    for m in _FIO_INITIALS_RE.finditer(text):
        # Пропускаем если уже покрыто полным ФИО
        overlap = any(s <= m.start() and m.end() <= e for s, e in seen_spans)
        if overlap:
            continue
        raw = m.group().strip()
        try:
            canonical = normalizer.normalize_value(raw)
        except Exception:
            canonical = raw
        matches.append(EntityMatch(
            entity_type="fio",
            original=raw,
            canonical=canonical,
            start=m.start(),
            end=m.end(),
            confidence=0.75,
        ))
    return matches


def _find_addresses(text: str) -> list[EntityMatch]:
    normalizer = AddressNormalizer()
    matches = []
    # Ищем предложения/фрагменты, содержащие адресные ключевые слова
    sentences = re.split(r"[.;\n]", text)
    pos = 0
    for sent in sentences:
        if _ADDRESS_KEYWORDS.search(sent):
            raw = sent.strip()
            if len(raw) > 5:
                try:
                    canonical = normalizer.normalize_value(raw)
                except Exception:
                    canonical = raw
                start = text.find(raw, pos)
                end = start + len(raw) if start != -1 else pos + len(raw)
                matches.append(EntityMatch(
                    entity_type="address",
                    original=raw,
                    canonical=canonical,
                    start=max(start, 0),
                    end=end,
                    confidence=0.8,
                ))
        pos += len(sent) + 1
    return matches


def _find_organizations(text: str) -> list[EntityMatch]:
    normalizer = OrganizationNormalizer()
    matches = []
    for m in _ORG_KEYWORDS.finditer(text):
        # Захватываем ОПФ + следующие 60 символов как название
        start = m.start()
        end = min(m.end() + 60, len(text))
        raw = text[start:end].split("\n")[0].strip()
        if len(raw) < 3:
            continue
        try:
            canonical = normalizer.normalize_value(raw)
        except Exception:
            canonical = raw
        matches.append(EntityMatch(
            entity_type="organization",
            original=raw,
            canonical=canonical,
            start=start,
            end=start + len(raw),
            confidence=0.8,
        ))
    return matches


# ---------------------------------------------------------------------------
# Основная функция нормализации документа
# ---------------------------------------------------------------------------

def normalize_document(
    text: str,
    enabled_types: list[str] | None = None,
) -> DocumentNormalizationResult:
    """
    Найти все PII-сущности в тексте и вернуть нормализованный текст + маппинг.

    enabled_types: список типов для поиска. По умолчанию все:
        ["fio", "phone", "email", "inn", "address", "organization"]
    """
    all_types = ["fio", "phone", "email", "inn", "address", "organization"]
    if enabled_types is None:
        enabled_types = all_types

    finders = {
        "fio": _find_fio,
        "phone": _find_phones,
        "email": _find_emails,
        "inn": _find_inns,
        "address": _find_addresses,
        "organization": _find_organizations,
    }

    all_entities: list[EntityMatch] = []
    for etype in enabled_types:
        if etype in finders:
            found = finders[etype](text)
            all_entities.extend(found)

    # Сортируем по позиции в тексте (с конца, чтобы замены не сбивали индексы)
    all_entities.sort(key=lambda e: e.start, reverse=True)

    # Убираем дубли (одна позиция — одна замена)
    seen_positions: set[int] = set()
    unique_entities: list[EntityMatch] = []
    for ent in all_entities:
        if ent.start not in seen_positions:
            seen_positions.add(ent.start)
            unique_entities.append(ent)

    # Применяем замены в тексте
    normalized_text = text
    for ent in unique_entities:
        if ent.original != ent.canonical:
            normalized_text = (
                normalized_text[: ent.start]
                + ent.canonical
                + normalized_text[ent.end :]
            )

    # Строим маппинг
    by_type: dict[str, list[dict]] = {t: [] for t in all_types}
    replacement_map: dict[str, str] = {}
    stats: dict[str, int] = {t: 0 for t in all_types}

    for ent in reversed(unique_entities):  # обратно в порядке документа
        by_type[ent.entity_type].append({
            "original": ent.original,
            "canonical": ent.canonical,
            "position": ent.start,
            "confidence": round(ent.confidence, 3),
            "meta": ent.meta,
        })
        replacement_map[ent.original] = ent.canonical
        stats[ent.entity_type] += 1

    mapping = {
        "meta": {
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "total_entities": len(unique_entities),
            "stats_by_type": stats,
        },
        "replacement_map": replacement_map,
        "by_type": by_type,
    }

    return DocumentNormalizationResult(
        normalized_text=normalized_text,
        mapping=mapping,
        entities=list(reversed(unique_entities)),
        stats=stats,
    )


# ---------------------------------------------------------------------------
# Нормализация DOCX с сохранением форматирования
# ---------------------------------------------------------------------------

def normalize_docx(content: bytes, enabled_types: list[str] | None = None) -> tuple[bytes, dict]:
    """
    Нормализует DOCX: заменяет найденные сущности прямо в параграфах документа.
    Возвращает (bytes нового DOCX, mapping dict).
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError("Установите python-docx: pip install python-docx")

    full_text, doc = extract_text_from_docx(content)
    result = normalize_document(full_text, enabled_types)

    # Применяем замены к каждому параграфу
    for para in doc.paragraphs:
        new_text = para.text
        for original, canonical in result.mapping["replacement_map"].items():
            if original in new_text:
                new_text = new_text.replace(original, canonical)
        if new_text != para.text:
            # Сохраняем первый run с новым текстом, остальные очищаем
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = new_text

    # Также обрабатываем таблицы в документе
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new_text = para.text
                    for original, canonical in result.mapping["replacement_map"].items():
                        if original in new_text:
                            new_text = new_text.replace(original, canonical)
                    if new_text != para.text and para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            run.text = ""

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), result.mapping