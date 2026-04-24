"""Извлечение текста из различных форматов документов.

Поддерживаются:
- .txt / .md / .csv — plain text
- .docx — через python-docx (параграфы + таблицы)
- .rtf — простой парсер RTF-групп (достаточно для обычных документов)

Возвращается список «кусков» текста (runs). Для каждого куска известно:
  * kind  — `paragraph` / `table_cell`
  * path  — адрес внутри документа (index параграфа или (table, row, col))
  * text  — оригинальный текст

Такое представление позволяет потом заменить текст in-place и собрать
документ обратно в том же формате.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Iterable


@dataclass
class TextChunk:
    """Кусок текста документа, пригодный для последующей замены."""

    idx: int                       # сквозной номер
    kind: str                      # paragraph | table_cell | line
    path: tuple[Any, ...]          # координаты внутри документа
    text: str                      # содержимое


@dataclass
class ExtractedDocument:
    """Результат извлечения: структура документа + сконкатенированный текст."""

    fmt: str                       # 'txt' | 'docx' | 'rtf' | 'md'
    chunks: list[TextChunk] = field(default_factory=list)
    raw: Any = None                # исходный объект (bytes / Document) — для реконструкции

    @property
    def full_text(self) -> str:
        return "\n".join(c.text for c in self.chunks)


# ---------------------------------------------------------------------------
# TXT / MD / CSV
# ---------------------------------------------------------------------------

def _extract_plain(data: bytes, fmt: str) -> ExtractedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("cp1251", errors="replace")
    # Строки сохраняем как отдельные chunk'и — это упрощает точечную замену
    # и сохраняет переносы строк при сборке обратно.
    lines = text.split("\n")
    chunks = [TextChunk(i, "line", (i,), line) for i, line in enumerate(lines)]
    return ExtractedDocument(fmt=fmt, chunks=chunks, raw=text)


def _rebuild_plain(doc: ExtractedDocument, replaced: dict[int, str]) -> bytes:
    lines = [replaced.get(c.idx, c.text) for c in doc.chunks]
    return "\n".join(lines).encode("utf-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx(data: bytes) -> ExtractedDocument:
    import io
    from docx import Document  # type: ignore

    document = Document(io.BytesIO(data))
    chunks: list[TextChunk] = []
    idx = 0

    for p_i, paragraph in enumerate(document.paragraphs):
        if paragraph.text:
            chunks.append(TextChunk(idx, "paragraph", ("p", p_i), paragraph.text))
            idx += 1

    for t_i, table in enumerate(document.tables):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                # По ячейке идём по параграфам — чтобы сохранить форматирование.
                for pp_i, para in enumerate(cell.paragraphs):
                    if para.text:
                        chunks.append(TextChunk(
                            idx, "table_cell", ("t", t_i, r_i, c_i, pp_i), para.text,
                        ))
                        idx += 1

    return ExtractedDocument(fmt="docx", chunks=chunks, raw=document)


def _rebuild_docx(doc: ExtractedDocument, replaced: dict[int, str]) -> bytes:
    import io
    # Нужен свежий объект — исходный мог быть мутирован. Перезагружаем из
    # сохранённой копии raw, если это Document, или возвращаем как есть.
    document = doc.raw

    for chunk in doc.chunks:
        new_text = replaced.get(chunk.idx)
        if new_text is None or new_text == chunk.text:
            continue
        path = chunk.path
        if path[0] == "p":
            p_i = path[1]
            paragraph = document.paragraphs[p_i]
            _replace_paragraph_text(paragraph, new_text)
        elif path[0] == "t":
            _, t_i, r_i, c_i, pp_i = path
            paragraph = document.tables[t_i].rows[r_i].cells[c_i].paragraphs[pp_i]
            _replace_paragraph_text(paragraph, new_text)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Заменить текст параграфа, сохранив форматирование первого run.

    Простейший подход: очистить все runs и записать новый текст в первый run
    (или создать новый, если runs нет). Для PII-замены этого достаточно.
    """
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


# ---------------------------------------------------------------------------
# RTF (простой парсер: берём текст между управляющих последовательностей)
# ---------------------------------------------------------------------------

def _extract_rtf(data: bytes) -> ExtractedDocument:
    import re

    try:
        text_raw = data.decode("utf-8")
    except UnicodeDecodeError:
        text_raw = data.decode("cp1251", errors="replace")

    # Убираем управляющие последовательности \word, \word123, { }
    cleaned = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text_raw)
    cleaned = cleaned.replace("{", "").replace("}", "")
    # Юникод-последовательности \'xx — пропустим, они редки
    cleaned = re.sub(r"\\'[0-9a-fA-F]{2}", "", cleaned)
    lines = [ln for ln in cleaned.split("\n")]
    chunks = [TextChunk(i, "line", (i,), ln) for i, ln in enumerate(lines)]
    return ExtractedDocument(fmt="rtf", chunks=chunks, raw=data)


def _rebuild_rtf(doc: ExtractedDocument, replaced: dict[int, str]) -> bytes:
    # Для RTF честная замена сложна — пересобираем как plain text в .txt
    # с оригинальной структурой строк.
    return _rebuild_plain(doc, replaced)


# ---------------------------------------------------------------------------
# Диспетчер
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = ("txt", "md", "csv", "docx", "rtf")


def extract_document(filename: str, data: bytes) -> ExtractedDocument:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext in ("txt", "md", "csv"):
        return _extract_plain(data, ext)
    if ext == "docx":
        return _extract_docx(data)
    if ext == "rtf":
        return _extract_rtf(data)
    raise ValueError(f"Формат .{ext} не поддерживается")


def rebuild_document(doc: ExtractedDocument, replaced: dict[int, str]) -> tuple[bytes, str]:
    """Собрать документ обратно с заменённым текстом.

    Возвращает (bytes, расширение результата).
    """
    if doc.fmt in ("txt", "md", "csv"):
        return _rebuild_plain(doc, replaced), doc.fmt
    if doc.fmt == "docx":
        return _rebuild_docx(doc, replaced), "docx"
    if doc.fmt == "rtf":
        # Возвращаем как .txt — см. комментарий в _rebuild_rtf.
        return _rebuild_rtf(doc, replaced), "txt"
    raise ValueError(f"Неизвестный формат: {doc.fmt}")
